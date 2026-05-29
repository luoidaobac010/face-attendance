import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Đặt màu nền cho một ô trong bảng."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def create_element(name):
    return OxmlElement(name)

def main():
    doc = Document()
    
    # Thiết lập lề cho trang
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Thiết lập kiểu chữ mặc định (Normal)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2B, 0x2B, 0x2B) # Màu xám đậm dễ đọc

    # Màu sắc chủ đạo (Navy Blue và Slate)
    NAVY = RGBColor(0x1B, 0x36, 0x5D)
    SLATE = RGBColor(0x5C, 0x76, 0x8D)
    
    # ------------------ TIÊU ĐỀ CHÍNH ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("TÀI LIỆU HƯỚNG DẪN BẢO VỆ ĐỒ ÁN")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(6)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("HỆ THỐNG ĐIỂM DANH KHUÔN MẶT BẰNG THUẬT TOÁN PCA + SVM")
    sub_run.font.size = Pt(14)
    sub_run.font.bold = True
    sub_run.font.color.rgb = SLATE
    subtitle_p.paragraph_format.space_after = Pt(30)

    # Thêm một đường kẻ ngang trang trí
    p_sep = doc.add_paragraph()
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep_run = p_sep.add_run("❖  ❖  ❖  ❖  ❖")
    sep_run.font.color.rgb = SLATE
    p_sep.paragraph_format.space_after = Pt(24)

    # ------------------ PHẦN 1: TỔNG QUAN ------------------
    h1 = doc.add_paragraph()
    r1 = h1.add_run("I. TỔNG QUAN VỀ DỰ ÁN & Ý NGHĨA ĐỀ TÀI")
    r1.font.size = Pt(14)
    r1.font.bold = True
    r1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Hệ thống điểm danh khuôn mặt (Face Attendance System) là một ứng dụng thực tiễn vượt trội, thay thế các phương thức điểm danh truyền thống (như gọi tên, ký giấy, quẹt thẻ) bằng công nghệ sinh trắc học hiện đại. Hệ thống tích hợp xử lý ảnh thời gian thực từ Webcam nhằm tối ưu hóa thời gian và ngăn chặn hoàn toàn việc gian lận điểm danh.")
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph("Dự án sử dụng mô hình kết hợp học máy truyền thống cực kỳ tối ưu:")
    p.paragraph_format.space_after = Pt(4)
    
    bullet1 = doc.add_paragraph(style='List Bullet')
    r = bullet1.add_run("Haar Cascade Classifier: ")
    r.font.bold = True
    bullet1.add_run("Phát hiện vùng chứa khuôn mặt trong luồng video trực tiếp từ Camera.")
    
    bullet2 = doc.add_paragraph(style='List Bullet')
    r = bullet2.add_run("PCA (Principal Component Analysis): ")
    r.font.bold = True
    bullet2.add_run("Giảm chiều dữ liệu từ ảnh thô 100x100 (10,000 chiều) xuống còn 50 chiều đặc trưng cốt lõi (Eigenfaces), giúp tăng tốc độ xử lý hàng trăm lần.")
    
    bullet3 = doc.add_paragraph(style='List Bullet')
    r = bullet3.add_run("SVM (Support Vector Machine): ")
    r.font.bold = True
    bullet3.add_run("Phân loại tuyến tính với độ chính xác cao để nhận diện chính xác danh tính và ghi nhận thông tin vào tệp lưu trữ Excel tự động.")

    # ------------------ PHẦN 2: KIẾN TRÚC VÀ PIPELINE ------------------
    h2 = doc.add_paragraph()
    r2 = h2.add_run("II. QUY TRÌNH HOẠT ĐỘNG CỦA HỆ THỐNG (PIPELINE)")
    r2.font.size = Pt(14)
    r2.font.bold = True
    r2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Hệ thống hoạt động khép kín qua 3 bước cốt lõi sau:")
    p.paragraph_format.space_after = Pt(6)

    # Tạo bảng quy trình
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Shading Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Bước'
    hdr_cells[1].text = 'Tên Quy Trình'
    hdr_cells[2].text = 'Chi Tiết Kỹ Thuật Thực Hiện'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "1B365D")

    row_data = [
        ("Bước 1", "Đăng ký (Collect)", "Webcam mở lên, tìm khuôn mặt bằng Haar Cascade. Cắt vùng chứa khuôn mặt, chuẩn hóa kích thước về 100x100 pixels ở hệ màu xám (GrayScale). Chụp đủ 50 ảnh thô lưu vào folder dataset/."),
        ("Bước 2", "Huấn luyện (Train)", "Đọc toàn bộ ảnh từ dataset, trải phẳng ảnh thành vector 10,000 chiều. Sử dụng SVD để phân rã ma trận dữ liệu ảnh, lấy ra 50 Eigenfaces quan trọng nhất (PCA). Chiếu dữ liệu ảnh lên 50 chiều này, rồi đưa vào bộ phân lớp SVM (Linear Kernel) để tối ưu ranh giới phân biệt giữa các người. Lưu model thành file model_svm.pkl."),
        ("Bước 3", "Nhận diện & Ghi danh (Recognize)", "Camera chạy trực tiếp. Phát hiện mặt, cắt ảnh, trải phẳng và chiếu vào không gian 50 chiều của PCA đã lưu. Đưa vào SVM dự đoán nhãn người dùng cùng xác suất tin cậy (Probability). Nếu độ tin cậy > 60%, tiến hành ghi nhận tên, ngày và giờ điểm danh hiện hành vào file Excel attendance.xlsx.")
    ]

    for idx, (step, name, desc) in enumerate(row_data):
        row_cells = table.rows[idx + 1].cells
        row_cells[0].text = step
        row_cells[1].text = name
        row_cells[2].text = desc
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row_cells[0], "F0F4F8" if idx % 2 == 0 else "FFFFFF")
        set_cell_background(row_cells[1], "F0F4F8" if idx % 2 == 0 else "FFFFFF")
        set_cell_background(row_cells[2], "F0F4F8" if idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # ------------------ PHẦN 3: PHÂN TÍCH THUẬT TOÁN ------------------
    h3 = doc.add_paragraph()
    r3 = h3.add_run("III. PHÂN TÍCH CHI TIẾT CÁC THUẬT TOÁN")
    r3.font.size = Pt(14)
    r3.font.bold = True
    r3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)

    # Haar Cascade
    p = doc.add_paragraph()
    r = p.add_run("1. Phát hiện khuôn mặt bằng Haar Cascade (OpenCV)")
    r.font.bold = True
    r.font.color.rgb = SLATE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph("Phương pháp dựa trên các đặc trưng Haar (chữ nhật sáng/tối đại diện cho các vùng mắt, sống mũi, má trên mặt). Sử dụng cấu trúc hình ảnh tích hợp (Integral Image) để tính toán nhanh các đặc trưng, kết hợp thuật toán Boosting (Adaboost) để lọc các đặc trưng yếu và xếp tầng các bộ phân loại phân cấp (Cascade) nhằm loại bỏ nhanh các vùng không chứa mặt, tối ưu hiệu năng chạy thời gian thực.")
    p.paragraph_format.line_spacing = 1.2

    # PCA/SVD
    p = doc.add_paragraph()
    r = p.add_run("2. Giảm chiều dữ liệu bằng Principal Component Analysis (PCA)")
    r.font.bold = True
    r.font.color.rgb = SLATE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph("Mỗi ảnh khuôn mặt kích thước 100x100 có tổng cộng 10,000 điểm ảnh (pixels). Chạy trực tiếp trên 10,000 chiều sẽ tốn rất nhiều tài nguyên và dễ gây Overfitting. PCA giải quyết vấn đề này bằng cách tìm ra các hướng (components) có phương sai lớn nhất trong không gian ảnh. Dự án sử dụng kỹ thuật Phân rã giá trị suy biến (SVD - Singular Value Decomposition) trực tiếp trên ma trận dữ liệu ảnh đã trừ đi ảnh trung bình (Mean Face). Điều này giúp nén ảnh thô về 50 chiều đặc trưng mà vẫn giữ được > 95% thông tin cấu trúc khuôn mặt ban đầu.")
    p.paragraph_format.line_spacing = 1.2

    # SVM
    p = doc.add_paragraph()
    r = p.add_run("3. Phân loại bằng Support Vector Machine (SVM)")
    r.font.bold = True
    r.font.color.rgb = SLATE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph("SVM là thuật toán phân lớp cực kỳ mạnh mẽ cho dữ liệu có số chiều cao. Sau khi PCA giảm số chiều của ảnh về 50, SVM tiến hành tìm kiếm một siêu phẳng phân lớp tối ưu (Optimal Hyperplane) có lề (Margin) rộng nhất để ngăn tách dữ liệu của các cá nhân khác nhau. Dự án sử dụng nhân tuyến tính (Linear Kernel) kết hợp tính năng xác suất (`probability=True`) của thư viện Scikit-learn để lấy độ tin cậy của dự đoán, chỉ ghi nhận điểm danh khi độ tin cậy đạt mức an toàn (> 60%).")
    p.paragraph_format.line_spacing = 1.2

    # ------------------ PHẦN 4: CHI TIẾT CÁC TỆP TRONG MÃ NGUỒN ------------------
    h4_file = doc.add_paragraph()
    r4_file = h4_file.add_run("IV. CHI TIẾT CÁC TỆP TRONG MÃ NGUỒN DỰ ÁN")
    r4_file.font.size = Pt(14)
    r4_file.font.bold = True
    r4_file.font.color.rgb = NAVY
    h4_file.paragraph_format.space_before = Pt(18)
    h4_file.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Để Hội đồng thấy được tính tổ chức và sự chặt chẽ của dự án, dưới đây là bảng phân tích chi tiết vai trò của từng tệp tin trong cấu trúc thư mục:")
    p.paragraph_format.space_after = Pt(6)

    # Bảng phân tích tệp
    table_files = doc.add_table(rows=8, cols=3)
    table_files.style = 'Light Shading Accent 1'
    
    hdr_cells_f = table_files.rows[0].cells
    hdr_cells_f[0].text = 'Tên Tệp Tin'
    hdr_cells_f[1].text = 'Định Dạng'
    hdr_cells_f[2].text = 'Chức Năng & Nhiệm Vụ Chi Tiết'
    
    for cell in hdr_cells_f:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(cell, "1B365D")

    files_data = [
        ("main.py", "Mã nguồn Python (.py)", "Trái tim điều khiển của cả hệ thống. File này cung cấp menu tương tác (C, R, T). Chứa toàn bộ các hàm cốt lõi: từ collect khuôn mặt, load dữ liệu thô, thực hiện nén giảm chiều ảnh (PCA), huấn luyện SVM, khởi động Webcam nhận diện điểm danh và tự động ghi nhận log qua mark_once."),
        ("collect.py", "Mã nguồn Python (.py)", "Module hỗ trợ phụ trách đăng ký. Khi chạy độc lập, tệp này mở camera, sử dụng Haar Cascade phát hiện mặt, cắt ảnh và tự động chuẩn hóa về kích thước 100x100 pixels ở chế độ ảnh xám (Grayscale), lưu đủ 50 ảnh vào dataset/."),
        ("train.py", "Mã nguồn Python (.py)", "Module huấn luyện độc lập tùy chọn. Cho phép duyệt qua dataset và huấn luyện mô hình dựa trên thuật toán LBPH Face Recognizer truyền thống của OpenCV để lưu trữ mô hình dạng yml, giúp so sánh hiệu năng và độ chính xác với mô hình SVM."),
        ("recognize.py", "Mã nguồn Python (.py)", "Module nhận diện độc lập tùy chọn. Sử dụng thuật toán LBPH Face Recognizer đã được lưu ở file trainer.yml để nhận dạng đối tượng nhanh mà không cần giảm chiều PCA."),
        ("attendance.xlsx", "Bảng tính Excel (.xlsx)", "Cơ sở dữ liệu lưu giữ kết quả điểm danh. Lưu thông tin gồm: Tên người dùng được nhận diện, Ngày điểm danh và Giờ điểm danh cụ thể dưới định dạng dòng cột của Pandas."),
        ("model_svm.pkl", "Tệp nhị phân (.pkl)", "Tệp lưu trữ trạng thái mô hình học máy sau khi train thành công thông qua thư viện pickle. Lưu ma trận chiếu PCA W, ảnh trung bình mean, trọng số phân lớp SVM và từ điển ánh xạ label_map."),
        ("features.npz", "Tệp nén dữ liệu (.npz)", "Lưu trữ nén mảng đặc trưng ảnh (faces) và nhãn (labels) tương ứng của Numpy. Giúp tối ưu hóa dung lượng ổ cứng, có thể xóa thư mục ảnh thô dataset/ đi mà vẫn có thể train lại bình thường bằng tệp đặc trưng này.")
    ]

    for idx, (fname, ftype, fdesc) in enumerate(files_data):
        row_cells = table_files.rows[idx + 1].cells
        row_cells[0].text = fname
        row_cells[1].text = ftype
        row_cells[2].text = fdesc
        row_cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_background(row_cells[0], "F0F4F8" if idx % 2 == 0 else "FFFFFF")
        set_cell_background(row_cells[1], "F0F4F8" if idx % 2 == 0 else "FFFFFF")
        set_cell_background(row_cells[2], "F0F4F8" if idx % 2 == 0 else "FFFFFF")

    doc.add_paragraph().paragraph_format.space_before = Pt(8)

    # ------------------ PHẦN 5: CÂU HỎI PHẢN BIỆN ------------------
    h4 = doc.add_paragraph()
    r4 = h4.add_run("V. BỘ CÂU HỎI PHẢN BIỆN THƯỜNG GẶP CỦA HỘI ĐỒNG & CÁCH TRẢ LỜI")
    r4.font.size = Pt(14)
    r4.font.bold = True
    r4.font.color.rgb = NAVY
    h4.paragraph_format.space_before = Pt(18)
    h4.paragraph_format.space_after = Pt(6)

    qa_list = [
        ("Câu 1: Tại sao em lại chọn mô hình PCA + SVM thay vì dùng các mô hình Deep Learning hiện đại (như CNN, ResNet)?",
         "Dạ thưa Thầy/Cô, việc chọn mô hình học máy truyền thống PCA + SVM mang lại 3 ưu thế vượt trội phù hợp với ngữ cảnh đồ án:\n"
         "1. Hiệu năng phần cứng: Mô hình PCA+SVM cực kỳ nhẹ, có thể huấn luyện và nhận diện mượt mà theo thời gian thực trực tiếp trên CPU của máy tính cá nhân hoặc các hệ thống nhúng (như Raspberry Pi) mà không cần trang bị card đồ họa GPU đắt tiền.\n"
         "2. Dữ liệu huấn luyện ít: Deep Learning cần hàng nghìn, hàng vạn bức ảnh của mỗi người để đạt độ chính xác tốt. Trong khi đó, với mô hình PCA+SVM này, mỗi người dùng mới đăng ký chỉ cần chụp khoảng 50 bức ảnh (mất chưa đầy 5 giây) là hệ thống đã có thể học và nhận diện vô cùng chính xác.\n"
         "3. Tốc độ huấn luyện siêu tốc: Việc huấn luyện lại mô hình (khi có người mới đăng ký) diễn ra gần như tức thời (chưa tới 1 giây), mang lại trải nghiệm người dùng tối ưu."),
        
        ("Câu 2: Tại sao em lại sử dụng ảnh xám (GrayScale) thay vì ảnh màu (RGB)?",
         "Dạ thưa Thầy/Cô, đặc trưng màu sắc của ảnh (đỏ, xanh, vàng...) phụ thuộc rất lớn vào điều kiện ánh sáng môi trường xung quanh (đèn neon, ánh sáng mặt trời) và màu da của đối tượng. Điều này dễ làm mô hình bị nhiễu. Bằng việc chuyển sang ảnh xám (Grayscale), chúng ta chỉ tập trung trích xuất cấu trúc hình học của khuôn mặt (khoảng cách giữa hai mắt, chiều cao mũi, góc cằm) và phân bổ mức độ sáng tối. Hơn thế nữa, ảnh xám giảm lượng dữ liệu lưu trữ xuống 3 lần so với ảnh màu (chỉ còn 1 kênh màu thay vì 3 kênh màu), giúp hệ thống chạy nhanh hơn đáng kể."),
        
        ("Câu 3: PCA giảm chiều bằng SVD hoạt động như thế nào trong chương trình của em?",
         "Dạ thưa Thầy/Cô, quy trình hoạt động của PCA được thực hiện tuần tự như sau:\n"
         "- Đầu tiên, tất cả các ảnh thô 100x100 được 'trải phẳng' thành một mảng vector 1 chiều có kích thước 10,000.\n"
         "- Tiếp theo, ta tính toán ảnh trung bình (Mean Face) của toàn bộ dữ liệu, rồi lấy từng ảnh trừ đi ảnh trung bình này để chuẩn hóa dữ liệu tập trung xung quanh gốc tọa độ.\n"
         "- Tiếp đó, ta áp dụng thuật toán SVD (Singular Value Decomposition) để phân rã ma trận dữ liệu thành ba ma trận thành phần U, S và Vt. Ma trận Vt chứa các vector riêng (chính là các trục biểu diễn có phương sai từ lớn đến bé).\n"
         "- Cuối cùng, em trích xuất lấy 50 vector riêng có trị riêng lớn nhất (Eigenfaces) để tạo thành ma trận chiếu W. Bất kỳ khuôn mặt mới nào chụp từ webcam chỉ cần nhân ma trận với ma trận chiếu này để thu gọn về 50 chỉ số đặc trưng."),
         
        ("Câu 4: Làm thế nào hệ thống của em giải quyết vấn đề điểm danh trùng lặp trong một ngày?",
         "Dạ thưa Thầy/Cô, tính năng này được xử lý triệt để thông qua hàm `mark_once(name)`.\n"
         "- Khi nhận diện thành công một người, hệ thống sẽ mở tệp lưu trữ Excel `attendance.xlsx` lên bằng thư viện Pandas.\n"
         "- Hệ thống sẽ lấy ngày hiện tại (ví dụ: 2026-05-20) và so sánh điều kiện đồng thời: tên trùng khớp VÀ ngày điểm danh trùng khớp ngày hôm nay.\n"
         "- Nếu đã có bản ghi trùng khớp với ngày hôm nay, hàm sẽ trả về `False` và bỏ qua không ghi thêm nữa để tránh trùng lặp.\n"
         "- Sang ngày tiếp theo, điều kiện ngày hôm nay thay đổi nên hệ thống sẽ tự động cho phép điểm danh lại bình thường mà không bị block vĩnh viễn. Ngoài ra, chương trình có hàm tự động sửa lỗi cấu trúc tệp Excel nếu bị thiếu hoặc lệch cột, đảm bảo tính bền vững của dữ liệu.")
    ]

    for q_idx, (q, a) in enumerate(qa_list):
        # Câu hỏi
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(10)
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(f"✦ {q}")
        r_q.font.bold = True
        r_q.font.color.rgb = NAVY
        
        # Câu trả lời
        p_a = doc.add_paragraph()
        p_a.paragraph_format.left_indent = Inches(0.25)
        p_a.paragraph_format.space_after = Pt(8)
        p_a.paragraph_format.line_spacing = 1.15
        p_a.add_run(a)

    # ------------------ PHẦN 6: HƯỚNG DẪN DEMO ------------------
    h5 = doc.add_paragraph()
    r5 = h5.add_run("VI. KỊCH BẢN DEMO THUYẾT PHỤC TRƯỚC HỘI ĐỒNG")
    r5.font.size = Pt(14)
    r5.font.bold = True
    r5.font.color.rgb = NAVY
    h5.paragraph_format.space_before = Pt(18)
    h5.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Để buổi bảo vệ đồ án diễn ra ấn tượng và mượt mà nhất, em nên thực hiện demo theo kịch bản 3 bước sau:")
    p.paragraph_format.space_after = Pt(6)

    demo_steps = [
        ("Bước 1: Giới thiệu giao diện và Train mẫu cũ",
         "Khởi động chương trình `python main.py`. Chọn chức năng **`C`** (Chấm công). Camera mở lên và ngay lập tức nhận diện thành công các mẫu có sẵn của người khác hoặc khuôn mặt em nếu đã được nạp sẵn. Nhấn **ESC** để tắt camera."),
        ("Bước 2: Demo đăng ký người mới cực nhanh",
         "Chọn chức năng **`R`** (Đăng ký người mới). Nhập tên một thành viên mới (hoặc tên Thầy/Cô chấm phản biện nếu Thầy/Cô muốn trải nghiệm thử). Camera mở lên chụp 50 ảnh thô trong 3-5 giây. Lúc này, nhấn mạnh với hội đồng là: *'Hệ thống đang chụp 50 ảnh ở các góc nghiêng khác nhau để tạo độ đa dạng cho tập dữ liệu học'*."),
        ("Bước 3: Huấn luyện tức thời & Nhận diện thành công",
         "Ngay sau khi chụp đủ 50 ảnh, hệ thống tự động xóa mô hình cũ, trích xuất đặc trưng PCA, huấn luyện lại SVM trong chớp mắt và bật camera nhận diện người vừa đăng ký thành công. Mở file `attendance.xlsx` lên cho Hội đồng xem dòng thông tin điểm danh của người mới vừa được tự động thêm vào với ngày giờ chính xác."),
    ]

    for step_title, step_desc in demo_steps:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(6)
        p_d.paragraph_format.space_after = Pt(4)
        r_d = p_d.add_run(f"✔ {step_title}")
        r_d.font.bold = True
        r_d.font.color.rgb = SLATE
        
        p_desc = doc.add_paragraph(step_desc)
        p_desc.paragraph_format.left_indent = Inches(0.2)
        p_desc.paragraph_format.space_after = Pt(6)
        p_desc.paragraph_format.line_spacing = 1.15

    # ------------------ KẾT LUẬN ------------------
    h6 = doc.add_paragraph()
    r6 = h6.add_run("VII. KẾT LUẬN & HƯỚNG PHÁT TRIỂN")
    r6.font.size = Pt(14)
    r6.font.bold = True
    r6.font.color.rgb = NAVY
    h6.paragraph_format.space_before = Pt(18)
    h6.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Đồ án đã xây dựng thành công một hệ thống điểm danh khuôn mặt hoạt động ổn định, chính xác cao và tối ưu tài nguyên vượt trội bằng việc kết hợp Haar Cascade, PCA và SVM. Hệ thống đáp ứng tốt các tiêu chí xử lý thời gian thực và lưu trữ dữ liệu an toàn.")
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph("Hướng phát triển trong tương lai:")
    p.paragraph_format.space_after = Pt(4)

    dev1 = doc.add_paragraph(style='List Bullet')
    dev1.add_run("Tích hợp chức năng chống giả mạo khuôn mặt (Liveness Detection) bằng cách bắt người dùng chớp mắt hoặc cười trước camera để ngăn chặn hành vi sử dụng ảnh chụp trên điện thoại để điểm danh hộ.")
    
    dev2 = doc.add_paragraph(style='List Bullet')
    dev2.add_run("Đưa dữ liệu điểm danh lưu trữ lên hệ thống cơ sở dữ liệu đám mây (MySQL/Firebase) kết hợp giao diện Web Dashboard quản lý trực quan cho quản trị viên.")

    # Lưu tài liệu thành file Word
    output_filename = "HUONG_DAN_BAO_VE_DU_AN_CHI_TIET.docx"
    doc.save(output_filename)
    print("Done! Word file created successfully: " + output_filename)

if __name__ == "__main__":
    main()
